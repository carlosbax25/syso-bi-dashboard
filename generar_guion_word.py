"""Genera el guion de sustentación en Word."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

SLIDES = [
    {
        'titulo': 'Slide 1 — PORTADA',
        'texto': 'Buenos días/tardes. Somos Luis Arrieta, Carlos Acosta y Pedro Guzmán, estudiantes de la Especialización en Analítica de Datos e Información del Tecnológico Comfenalco. Hoy presentamos nuestro proyecto integrador: el diseño de una solución de inteligencia de negocios para la gestión de órdenes de servicio ARL en una IPS del sector de seguridad y salud en el trabajo.\n\nNota: Este es un proyecto académico con datos sintéticos y una organización representativa con fines ilustrativos.',
    },
    {
        'titulo': 'Slide 2 — APERTURA (¿Por qué importa?)',
        'texto': '¿Por qué es importante este proyecto? Porque aborda un problema real del sector salud: la facturación. En una IPS que presta servicios a las Administradoras de Riesgos Laborales, la facturación es la principal fuente de ingresos. Cada día que pasa sin facturar un servicio ya ejecutado es un día de riesgo financiero.\n\nLos números hablan por sí solos: estamos hablando de 45.000 órdenes de servicio gestionadas manualmente en hojas de cálculo, $926 millones de pesos en cartera pendiente por facturar, un promedio de 22.7 días entre ejecutar un servicio y emitir la factura, y cero visibilidad en tiempo real del estado de esas órdenes.\n\nLa pregunta que nos hicimos fue: ¿cómo transformamos esos datos dispersos en información que permita tomar decisiones oportunas?',
    },
    {
        'titulo': 'Slide 3 — CONTEXTO (¿Dónde ocurre?)',
        'texto': 'El contexto de este proyecto es una IPS ubicada en Cartagena, Colombia, dedicada a la seguridad y salud en el trabajo. Esta organización presta servicios a 4 Administradoras de Riesgos Laborales: Sura, Colmena, Bolívar y Alfa.\n\nSu portafolio incluye 76 servicios diferentes, clasificados en 4 tipos: laboratorio clínico, pruebas complementarias, asesoría y capacitación. Atiende a 14 empresas clientes y ha acumulado 45.000 órdenes de servicio en el periodo 2020 a 2026.\n\nToda esta información se gestiona mediante la Matriz de Gestión ARL, que es un archivo Excel donde se registran manualmente todas las órdenes desde que se reciben hasta que se facturan.',
    },
    {
        'titulo': 'Slide 4 — PROBLEMA (¿Qué situación se quiere resolver?)',
        'texto': 'El problema central que identificamos es la gestión ineficiente de las órdenes de servicio. Las causas raíz son cuatro: primero, los registros se llevan manualmente en Excel sin ninguna automatización. Segundo, no hay integración entre los diferentes sistemas de información. Tercero, no existe un modelo de datos estructurado. Y cuarto, los registros presentan baja calidad y estandarización.\n\n¿Cuáles son las consecuencias? El tiempo promedio para facturar un servicio ya ejecutado es de 22.7 días. Hay 18.912 órdenes que ya fueron ejecutadas pero que aún no tienen factura emitida, representando $926 millones en cartera pendiente. Esto genera reprocesos, sobrecarga laboral y, lo más crítico, decisiones que se toman sin respaldo analítico.',
    },
    {
        'titulo': 'Slide 5 — OBJETIVO (¿Qué se propuso lograr?)',
        'texto': 'Nuestro objetivo general fue diseñar una solución de analítica de datos que permita integrar, transformar y visualizar la información de las órdenes ARL mediante un tablero de inteligencia de negocios desarrollado en Python.\n\nPara lograrlo definimos 5 objetivos específicos: primero, identificar las fuentes de datos y las variables clave del proceso. Segundo, diseñar un modelo unificado de datos. Tercero, diseñar el flujo de extracción, transformación y carga de la información. Cuarto, definir los indicadores clave de desempeño. Y quinto, proponer el diseño funcional del tablero interactivo.\n\nCada uno de estos objetivos es clickeable en la presentación y pueden ver el detalle de lo que logramos en cada uno.',
    },
    {
        'titulo': 'Slide 6 — METODOLOGÍA (¿Cómo se desarrolló?)',
        'texto': 'Seguimos la metodología CRISP-DM, que es el estándar para proyectos de analítica de datos. Consta de 6 fases que ejecutamos de la siguiente manera:\n\nEn la fase de Negocio realizamos entrevistas con el personal administrativo para diagnosticar el problema. En la fase de Datos revisamos la Matriz ARL completa e identificamos 13 variables relevantes. En la fase de Preparación construimos un proceso automatizado de carga desde Excel a la aplicación. En la fase de Modelado implementamos 6 indicadores de desempeño y una proyección de ingresos a 6 meses. En la fase de Evaluación validamos los indicadores con el personal de la IPS y ajustamos los umbrales de alerta. Y en la fase de Despliegue publicamos el tablero en internet, accesible desde cualquier navegador.',
    },
    {
        'titulo': 'Slide 7 — SOLUCIÓN (¿Qué se construyó?)',
        'texto': '¿Qué construimos? Un tablero de indicadores completo con dos grandes componentes.\n\nEl primero es el tablero principal, que incluye 6 indicadores interactivos — al hacer clic en cada uno se abre el detalle. Tiene 4 gráficos dinámicos con filtros, filtros combinados por ARL, tipo de servicio, programa y tarea, y permite exportar los datos filtrados a Excel.\n\nEl segundo componente son las funcionalidades de análisis: una proyección de ingresos a 6 meses basada en la tendencia histórica, un resumen ejecutivo narrativo que se genera automáticamente, un análisis de la cartera pendiente y los tiempos de facturación, la posibilidad de cargar datos masivamente desde un archivo Excel, y un diseño accesible para personas con daltonismo.\n\nTodo esto fue desarrollado en Python y está publicado en internet, accesible desde cualquier navegador sin necesidad de instalar nada.',
    },
    {
        'titulo': 'Slide 8 — RESULTADOS (¿Qué se obtuvo?)',
        'texto': 'Los resultados se resumen en estos 6 indicadores calculados sobre las 45.000 órdenes del periodo 2020 a 2026:\n\nEl total de órdenes procesadas es de 45.000. Los ingresos totales suman $2.204 millones de pesos. La tasa de cumplimiento — es decir, el porcentaje de órdenes que avanzan a estados completados — es del 62.2%. Hay 11.633 órdenes actualmente en gestión activa. La cartera sin facturar asciende a $926 millones. Y el tiempo promedio entre ejecutar un servicio y facturarlo es de 22.7 días.\n\nEl hallazgo más importante es que el 42% de las órdenes que ya fueron completadas operativamente no han sido facturadas. Esto nos dice que el punto crítico del proceso no está en prestar el servicio — eso se hace bien — sino en facturarlo después.',
    },
    {
        'titulo': 'Slide 9 — VALIDACIÓN (¿Cómo se comprobó?)',
        'texto': 'La solución se validó en 5 niveles. A nivel técnico, creamos 23 pruebas automatizadas que verifican el correcto funcionamiento de todos los indicadores y filtros. A nivel del modelo predictivo, reservamos los últimos 3 meses de datos como prueba y comparamos las proyecciones contra los valores reales, obteniendo métricas aceptables. A nivel funcional, realizamos sesiones presenciales con el Analista de Gestión ARL y con el Director Administrativo, quienes confirmaron que los indicadores son coherentes con los datos que manejan manualmente. Y a nivel normativo, contrastamos el umbral de 30 días con los plazos contractuales reales de las ARL.\n\nPueden hacer clic en cada fila de la tabla para ver el detalle de cómo se comprobó cada nivel.',
    },
    {
        'titulo': 'Slide 10 — VALOR GENERADO (¿Qué aporta?)',
        'texto': '¿Qué valor genera esta solución? Lo resumimos en 4 indicadores de impacto medibles.\n\nPrimero, el tiempo de generación de reportes pasó de 3 a 4 horas a segundos. Segundo, la detección de órdenes sin facturar, que antes dependía de una revisión manual semanal, ahora es permanente y automática. Tercero, la visibilidad del estado de las órdenes, que antes solo estaba disponible al cierre del mes, ahora es visible en todo momento. Y cuarto, la proyección de ingresos no existía antes — ahora se pueden proyectar 6 meses por cada ARL.\n\nAdemás, la arquitectura fue diseñada para crecer: está preparada para conectarse a una base de datos permanente, su información puede consumirse desde otros sistemas, y puede replicarse en otras instituciones del sector.',
    },
    {
        'titulo': 'Slide 11 — CONCLUSIONES (¿Qué se aprendió?)',
        'texto': 'Llegamos a 4 conclusiones principales.\n\nPrimera: logramos implementar una solución de inteligencia de negocios funcional con 6 indicadores y proyección predictiva, desplegada en un entorno de producción real accesible desde internet.\n\nSegunda: identificamos que el punto crítico del proceso está en facturar, no en prestar el servicio. Las órdenes se ejecutan, pero se quedan estancadas en la etapa administrativa.\n\nTercera: la herramienta reduce el tiempo de generación de reportes de varias horas a consulta inmediata, habilitando decisiones oportunas.\n\nY cuarta: demostramos que es viable transformar la gestión de una IPS mediante analítica de datos usando herramientas de código abierto y sin costo de infraestructura.',
    },
    {
        'titulo': 'Slide 12 — RECOMENDACIONES (¿Qué sigue?)',
        'texto': 'Para el futuro del proyecto recomendamos acciones en tres horizontes.\n\nA corto plazo: migrar los datos a una base de datos permanente, establecer conexión automática con los portales de las ARL para que las órdenes se carguen sin intervención manual, e implementar alertas por correo electrónico cuando una orden se acerque al umbral de riesgo.\n\nA mediano plazo: incorporar proyecciones que consideren los meses altos y bajos del año, implementar alertas automáticas para órdenes en riesgo de ser rechazadas, y que el propio tablero sugiera acciones específicas.\n\nA largo plazo: replicar esta solución en otras IPS del sector, implementar trazabilidad completa del proceso de principio a fin, y conectar el sistema con la facturación electrónica de la DIAN.',
    },
    {
        'titulo': 'Slide 13 — DEMOSTRACIÓN',
        'texto': '"Ahora les vamos a mostrar el tablero en funcionamiento. Contiene 6 indicadores interactivos, 4 gráficos y procesa 45.000 órdenes."\n\nAcciones durante la demo:\n1. Hacer clic en "Abrir tablero"\n2. Ingresar credenciales\n3. Mostrar las 6 tarjetas de indicadores — hacer clic en una para mostrar el detalle\n4. Mostrar los filtros en cascada — filtrar por una ARL específica\n5. Mostrar la proyección de ingresos (+6 meses)\n6. Ir al módulo de Resumen Ejecutivo\n7. Ir al módulo de Gestión y mostrar la carga de datos\n8. Volver al tablero principal\n\nTiempo sugerido: 3-5 minutos de demo en vivo.',
    },
    {
        'titulo': 'Slide 14 — GRACIAS / PREGUNTAS',
        'texto': '"Eso es todo por nuestra parte. Agradecemos su atención y estamos disponibles para responder preguntas."',
    },
]

PREGUNTAS = [
    ('¿Por qué Python y no Power BI?', 'Porque queríamos control total sobre la lógica de negocio, personalización sin límites de licencia, y despliegue gratuito en internet. Power BI requiere licencia y no permite el nivel de interactividad que logramos. Además, Python nos permite implementar modelos predictivos directamente dentro de la misma aplicación sin depender de herramientas externas.'),
    ('¿Los datos son reales?', 'Los datos son sintéticos pero basados en el catálogo real de servicios de una IPS. Los 76 servicios, las 4 ARLs y los valores del tarifario son reales; las combinaciones, fechas y volúmenes son generados algorítmicamente con una semilla reproducible para garantizar consistencia entre demostraciones.'),
    ('¿Cómo se garantiza la calidad de los datos?', 'El proceso de carga valida campo a campo: verifica tipos de dato, convierte fechas al formato estándar, asigna valores por defecto cuando hay campos vacíos, y reporta las filas con errores sin detener la carga completa. Esto garantiza que los datos que llegan al tablero son coherentes y utilizables.'),
    ('¿Qué pasa si la IPS quiere usar datos reales?', 'Solo tiene que cargar su archivo Excel real desde el Panel de Gestión. La arquitectura fue diseñada con una capa de abstracción que permite cambiar la fuente de datos — pasar de Excel a una base de datos como MySQL o PostgreSQL — sin modificar los indicadores, los gráficos ni el tablero.'),
    ('¿Por qué el umbral de 30 días?', 'Se validó con los plazos contractuales de las 4 ARLs. Superar ese tiempo expone a la IPS a glosas (rechazos de facturación) o pérdida del derecho a cobro. El umbral de 30 días es preventivo: da margen antes de llegar al límite contractual más estricto.'),
    ('¿Qué metodología usaron?', 'CRISP-DM, el estándar internacional para proyectos de analítica de datos, desarrollado por un consorcio de empresas líderes. Tiene 6 fases: comprensión del negocio, comprensión de los datos, preparación, modelado, evaluación y despliegue. Lo aplicamos secuencialmente durante las 14 semanas del semestre.'),
    ('¿Cómo funciona la proyección de ingresos?', 'Se toma el histórico de ingresos mensuales por cada ARL (últimos 12 meses completos), se calcula la tendencia mediante regresión lineal, y se proyecta esa tendencia 6 meses hacia el futuro. Es un modelo simple pero interpretable: la pendiente indica directamente si los ingresos de esa ARL están creciendo o decreciendo, y en qué magnitud mensual.'),
    ('¿Qué es la tasa de cumplimiento?', 'Es el porcentaje de órdenes que han avanzado a un estado de completado dentro del flujo. Incluye los estados Ejecutada, Soportes Radicados y Facturada. Si la tasa es del 62%, significa que el 62% de las órdenes recibidas ya pasaron por lo menos a la etapa de ejecución del servicio.'),
    ('¿Cómo se diferencia este proyecto de un simple reporte en Excel?', 'En Excel el reporte es estático: alguien lo genera manualmente cada semana o cada mes. Nuestro tablero es dinámico: los indicadores se calculan en tiempo real, se pueden filtrar por múltiples criterios simultáneamente, los gráficos se actualizan automáticamente, y hay alertas visuales cuando algo está fuera del rango esperado. Además, incluye capacidad predictiva que Excel no tiene.'),
    ('¿Cuánto cuesta mantener esta solución?', 'Actualmente cero pesos. Usamos herramientas de código abierto (Python, Flask) y hosting gratuito (Render.com). El único costo sería si se requiere una base de datos dedicada o un plan de hosting con mayor capacidad, pero para el volumen actual es suficiente la infraestructura gratuita.'),
    ('¿Qué limitaciones tiene el modelo predictivo?', 'Al ser regresión lineal simple, no captura estacionalidad (meses altos y bajos) ni eventos atípicos. Funciona bien para tendencias de corto plazo (6 meses) con datos estables. Para horizontes más largos o mayor precisión, se recomienda evolucionar a modelos de series de tiempo que consideren la estacionalidad.'),
    ('¿Cómo protegen los datos personales?', 'En el entorno público los datos son 100% sintéticos — no hay información real de ninguna persona o empresa. En producción, el tablero requiere autenticación obligatoria, la sesión se limpia al acceder al login, y no se almacenan datos en disco de forma permanente. Todo cumple con los principios de la Ley 1581 de 2012.'),
    ('¿Se puede replicar en otra IPS?', 'Sí. La arquitectura es genérica: el modelo de datos (13 campos + 10 estados) aplica a cualquier IPS que gestione órdenes ARL. Lo único específico es el catálogo de servicios, que es configurable. Una nueva IPS solo necesita cargar su propio archivo Excel para que el tablero funcione con sus datos.'),
    ('¿Qué aprendieron del proyecto?', 'Tres cosas principales: que el problema de facturación en IPS es más administrativo que operativo, que la analítica descriptiva bien aplicada genera valor inmediato sin necesidad de algoritmos complejos, y que es posible construir una solución profesional con herramientas gratuitas si se diseña la arquitectura correctamente desde el inicio.'),
    ('¿Por qué no usaron inteligencia artificial o machine learning avanzado?', 'Porque el volumen y la naturaleza de los datos no lo justifican. Con 12 puntos temporales por serie (meses), un modelo complejo se sobreajustaría. La regresión lineal es suficiente, interpretable por el usuario final, y no requiere infraestructura adicional. El valor del proyecto está en la integración y visualización de los datos, no en la complejidad del algoritmo.'),
]


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    for s in doc.sections:
        s.top_margin = Cm(2.54)
        s.bottom_margin = Cm(2.54)
        s.left_margin = Cm(2.54)
        s.right_margin = Cm(2.54)

    h = doc.add_heading('Guion de Sustentación', level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Especialización en Analítica de Datos · Tecnológico Comfenalco · 2026').font.color.rgb = RGBColor(0x5A, 0x6B, 0x7D)
    doc.add_paragraph()
    doc.add_paragraph('Tiempo total estimado: 20-25 minutos (incluyendo demo)')
    doc.add_paragraph('Tiempo sugerido por slide: 1.5-2 minutos')
    doc.add_paragraph()

    for slide in SLIDES:
        doc.add_heading(slide['titulo'], level=2)
        for para_text in slide['texto'].split('\n\n'):
            doc.add_paragraph(para_text)
        doc.add_paragraph()

    doc.add_page_break()
    doc.add_heading('Preguntas Frecuentes — Respuestas Preparadas', level=2)
    for pregunta, respuesta in PREGUNTAS:
        p = doc.add_paragraph()
        p.add_run(pregunta).bold = True
        doc.add_paragraph(respuesta)
        doc.add_paragraph()

    out = Path('static/Guion_Sustentacion.docx')
    doc.save(str(out))
    print(f'Generado: {out} ({out.stat().st_size/1024:.1f} KB)')


if __name__ == '__main__':
    main()
